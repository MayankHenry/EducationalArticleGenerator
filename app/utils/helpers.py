# app/utils/helpers.py
from io import BytesIO
from markdown_pdf import MarkdownPdf, Section
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown

def convert_markdown_to_pdf(md_content: str) -> BytesIO:
    """
    Converts a Markdown string to a PDF file in-memory.
    """
    pdf = MarkdownPdf()
    # Add a section to the PDF with the Markdown content
    pdf.add_section(Section(md_content))
    
    # Create an in-memory binary stream to store the PDF data
    buffer = BytesIO()
    pdf.save(buffer)
    
    # Rewind the buffer to the beginning so it can be read
    buffer.seek(0)
    return buffer

def convert_markdown_to_docx(md_content: str) -> BytesIO:
    """
    Converts a Markdown string to a DOCX file in-memory.
    This is a basic implementation that handles headings and paragraphs.
    """
    document = Document()
    
    # Split the content by lines to process headings and paragraphs
    for line in md_content.splitlines():
        line = line.strip()
        if not line:
            continue
        
        # Check for Markdown headings and apply the corresponding style
        if line.startswith("### "):
            document.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            document.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            document.add_heading(line[2:], level=1)
        else:
            # All other lines are treated as paragraphs
            document.add_paragraph(line)
            
    # Create an in-memory binary stream to store the DOCX data
    buffer = BytesIO()
    document.save(buffer)
    
    # Rewind the buffer to the beginning
    buffer.seek(0)
    return buffer